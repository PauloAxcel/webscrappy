import requests
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from urllib.parse import urljoin, urlparse
import time

def get_with_retries(url, max_retries=5, base_delay=30):
    """
    Attempts to fetch a URL with a robust retry mechanism and exponential backoff,
    especially useful for handling rate-limiting from servers.
    """
    for attempt in range(max_retries):
        try:
            response = requests.get(url, timeout=60)
            response.raise_for_status()  # Raise an HTTPError for bad responses (4xx or 5xx)
            return response
        except requests.exceptions.RequestException as e:
            if "10061" in str(e):
                print(f"🛑 Connection actively refused for {url}. Server is rate-limiting.")
            else:
                print(f"⚠️ Request failed for {url}: {e}")

            if attempt < max_retries - 1:
                delay = base_delay * (2 ** attempt)
                print(f"   -> Waiting for {delay} seconds before retrying...")
                time.sleep(delay)
            else:
                print(f"   -> Max retries reached for {url}. Giving up on this URL.")
                return None
    return None

def is_valid_url(url, original_site_domain):
    """
    Checks if the URL is a valid, followable link within the target Wayback Machine archive.
    """
    return (
        url.startswith("https://web.archive.org/web/") and
        original_site_domain in url and
        "mailto:" not in url and
        "javascript:" not in url
    )

def add_content_with_formatting(html_content, document):
    """
    Parses HTML content tag-by-tag to preserve basic formatting like headings,
    paragraphs, lists, and bold/italic text in the Word document.
    """
    if not html_content:
        return

    # Iterate through the main structural tags in the content area
    for tag in html_content.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol'], recursive=False):
        if tag.name in ['h1', 'h2', 'h3', 'h4']:
            level = int(tag.name[1])  # Gets the number from 'h1', 'h2', etc.
            document.add_heading(tag.get_text(strip=True), level=level)

        elif tag.name == 'p':
            p = document.add_paragraph()
            # This handles bold (<strong>, <b>) and italic (<em>, <i>) text within a paragraph
            for child in tag.children:
                if isinstance(child, NavigableString):
                    p.add_run(str(child))
                elif child.name in ['strong', 'b']:
                    p.add_run(child.get_text(strip=True)).bold = True
                elif child.name in ['em', 'i']:
                    p.add_run(child.get_text(strip=True)).italic = True

        elif tag.name in ['ul', 'ol']:
            # Handle bullet and numbered lists
            for li in tag.find_all('li', recursive=False):
                style = 'List Bullet' if tag.name == 'ul' else 'List Number'
                document.add_paragraph(li.get_text(strip=True), style=style)
        
        # Add a little space after each element for readability
        document.add_paragraph()

def scrape_page(url, document, visited_urls, processed_original_urls, unique_text_content, original_site_domain, output_filename):
    """
    Recursively scrapes a single page, handling duplicates, formatting content,
    and saving progress.
    """
    # 1. --- Duplicate Checking ---
    # This first check stops loops on the *exact same* Wayback URL
    if url in visited_urls:
        return

    # This smarter check stops re-scraping pages from different timestamps
    try:
        original_url = "http" + url.split("http", 2)[-1]
        if original_url in processed_original_urls:
            print(f"✅ Skipping already processed original URL: {original_url}")
            return
    except IndexError:
        print(f"⚠️ Could not extract original URL from: {url}")
        return

    print(f"\nAttempting to scrape: {url}")

    # 2. --- Fetching the Page ---
    response = get_with_retries(url)
    if not response:
        visited_urls.add(url) # Mark as visited even if it fails to avoid re-trying later
        return

    # Mark this page as fully processed
    visited_urls.add(url)
    processed_original_urls.add(original_url)
    time.sleep(1) # A polite 1-second delay between successful requests

    # 3. --- Parsing and Adding Content ---
    soup = BeautifulSoup(response.content, 'html.parser')
    for div in soup.find_all("div", {"id": "wm-ipp-base"}):
        div.decompose() # Remove the Wayback Machine's own header
    content_area = soup.find("body")

    # Use the raw HTML of the body as a key to check for unique content
    content_html = str(content_area)
    if content_html not in unique_text_content:
        unique_text_content.add(content_html)
        document.add_heading(f"Content from: {url}", level=1)
        
        # Call the smart formatting function
        add_content_with_formatting(content_area, document)

        document.add_page_break()
        
        # 4. --- Saving Progress ---
        try:
            document.save(output_filename)
            print(f"   -> ✅ Formatted progress saved to '{output_filename}'")
        except Exception as e:
            print(f"   -> ❌ Could not save progress: {e}")

    # 5. --- Finding and Following Links ---
    for link in content_area.find_all('a', href=True):
        href = link['href']
        absolute_url = urljoin(url, href)
        clean_url = urlparse(absolute_url)._replace(fragment="", query="", params="").geturl()

        if is_valid_url(clean_url, original_site_domain):
            # Recursively call the function for the new link
            scrape_page(clean_url, document, visited_urls, processed_original_urls, unique_text_content, original_site_domain, output_filename)

def main():
    """
    Main function to set up and start the scraping process.
    """
    start_url = "https://web.archive.org/web/20051027030857/http://maven.smith.edu/~thiebaut/ArtOfAssembly/artofasm.html"
    output_filename = "Formatted_Scraped_Content.docx"
    original_site_domain = "maven.smith.edu/~thiebaut/ArtOfAssembly"

    # --- Setup ---
    visited_urls = set()
    unique_text_content = set()
    processed_original_urls = set()
    document = Document()

    print("--- Starting Web Scraper ---")
    scrape_page(start_url, document, visited_urls, processed_original_urls, unique_text_content, original_site_domain, output_filename)
    print(f"\n🎉 Scraping process finished. Final content saved to '{output_filename}'")

if __name__ == "__main__":
    main()